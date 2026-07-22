import streamlit as st
import pandas as pd
from supabase import create_client, Client
from datetime import datetime
import io
import zipfile
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from PIL import Image

# --- PAGE CONFIGURATION ---
st.set_page_config(
    page_title="KEA Comprehensive School Management System",
    page_icon="🏫",
    layout="wide",
    initial_sidebar_state="expanded"
)

# --- SUPABASE CONNECTION SETUP ---
@st.cache_resource
def init_supabase():
    try:
        url = st.secrets["supabase"]["url"]
        key = st.secrets["supabase"]["key"]
        return create_client(url, key)
    except Exception as e:
        st.error(f"Failed to connect to Supabase. Check secrets.toml configuration. Details: {e}")
        return None

supabase = init_supabase()

# --- INITIALIZE SESSION STATE ---
if "logged_in" not in st.session_state:
    st.session_state.logged_in = False
    st.session_state.username = ""
    st.session_state.role = ""
    st.session_state.full_name = ""

# --- TEACHER SUBJECT & GRADE ASSIGNMENTS ---
TEACHER_ASSIGNMENTS = {
    "Eliars": {
        "assignments": [
            {"grade": "Grade 7", "subjects": ["ENGLISH", "PRETECHNICAL STUDIES"]},
            {"grade": "Grade 8", "subjects": ["ENGLISH", "C.A.S"]},
            {"grade": "Grade 9", "subjects": ["ENGLISH"]}
        ]
    },
    "Lucas": {
        "assignments": [
            {"grade": "Grade 7", "subjects": ["INTEGRATED SCIENCE"]},
            {"grade": "Grade 8", "subjects": ["INTEGRATED SCIENCE", "AGRICULTURE"]},
            {"grade": "Grade 9", "subjects": ["AGRICULTURE"]} # Class Teacher Grade 9
        ]
    },
    "Vincent": {
        "assignments": [
            {"grade": "Grade 7", "subjects": ["SOCIAL STUDIES"]},
            {"grade": "Grade 8", "subjects": ["SOCIAL STUDIES", "KISWAHILI"]}, # Class Teacher Grade 8
            {"grade": "Grade 9", "subjects": ["SOCIAL STUDIES", "KISWAHILI"]}
        ]
    },
    "Grace": {
        "assignments": [
            {"grade": "Grade 7", "subjects": ["RELIGIOUS EDUCATION (C.R.E)", "AGRICULTURE"]}, # Class Teacher Grade 7
            {"grade": "Grade 8", "subjects": ["RELIGIOUS EDUCATION (C.R.E)"]},
            {"grade": "Grade 9", "subjects": ["INTEGRATED SCIENCE"]}
        ]
    },
    "Elias": {
        "assignments": [
            {"grade": "Grade 7", "subjects": ["KISWAHILI"]}, # HOI
            {"grade": "Grade 9", "subjects": ["RELIGIOUS EDUCATION (C.R.E)"]}
        ]
    },
    "Valentine": {
        "assignments": [
            {"grade": "Grade 7", "subjects": ["MATHEMATICS"]},
            {"grade": "Grade 9", "subjects": ["MATHEMATICS"]}
        ]
    },
    "Elijah": {
        "assignments": [
            {"grade": "Grade 7", "subjects": ["C.A.S"]},
            {"grade": "Grade 8", "subjects": ["MATHEMATICS", "PRETECHNICAL STUDIES"]},
            {"grade": "Grade 9", "subjects": ["PRETECHNICAL STUDIES", "C.A.S"]}
        ]
    }
}

LEARNING_AREAS = [
    "MATHEMATICS", "ENGLISH", "KISWAHILI", "INTEGRATED SCIENCE", 
    "AGRICULTURE", "PRETECHNICAL STUDIES", "SOCIAL STUDIES", 
    "RELIGIOUS EDUCATION (C.R.E)", "C.A.S"
]

# --- AUTHENTICATION & LOGIN SCREEN ---
def login_screen():
    st.markdown("<h1 style='text-align: center; color: #1E3A8A;'>KEA COMPREHENSIVE SCHOOL</h1>", unsafe_allow_html=True)
    st.markdown("<h3 style='text-align: center; color: #4B5563;'>School Management & Information System</h3>", unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        if os.path.exists("logo.png"):
            st.image("logo.png", width=150)
        
        st.markdown("### Please Login to Continue")
        username = st.text_input("Username").strip()
        password = st.text_input("Password", type="password").strip()
        
        if st.button("Login", type="primary", use_container_width=True):
            if username == "Admin" and password == "janabi@26!":
                st.session_state.logged_in = True
                st.session_state.username = "Admin"
                st.session_state.role = "DEVELOPER"
                st.session_state.full_name = "System Developer"
                st.success("Welcome back, System Developer!")
                st.rerun()
                
            elif supabase:
                try:
                    res = supabase.table("teachers").select("*").ilike("username", username).execute()
                    if res.data:
                        user = res.data[0]
                        if password == "kea@26" or password == user.get("password"):
                            st.session_state.logged_in = True
                            st.session_state.username = user["username"]
                            st.session_state.role = user["designation"]
                            st.session_state.full_name = user["full_name"]
                            st.success(f"Welcome back, {user['full_name']}!")
                            st.rerun()
                        else:
                            st.error("Incorrect password. Default teacher password is 'kea@26'.")
                    else:
                        st.error("Username not found in teachers directory.")
                except Exception as e:
                    st.error(f"Database error during login: {e}")
            else:
                st.error("Supabase connection not initialized.")

if not st.session_state.logged_in:
    login_screen()
    st.stop()

# --- ROLE & NAVIGATION SETUP ---
role = st.session_state.role
is_admin_role = role in ["HOI", "DHOI", "DEVELOPER"]

st.sidebar.markdown(f"**Logged in as:** {st.session_state.full_name}")
st.sidebar.markdown(f"**Role:** {role}")
if st.sidebar.button("Logout"):
    st.session_state.logged_in = False
    st.rerun()

st.sidebar.markdown("---")
st.sidebar.markdown("### Navigation")

# RESTRICT NAVIGATION FOR NON-ADMIN TEACHERS
if is_admin_role:
    nav_options = [
        "Dashboard", 
        "Students Registration", 
        "Marks Entry", 
        "Results Analysis", 
        "Fee Payment", 
        "Teachers Portal", 
        "Teacher Time Login", 
        "Newsletter", 
        "School Contact"
    ]
else:
    nav_options = [
        "Dashboard", 
        "Marks Entry", 
        "Results Analysis", 
        "Teachers Portal", 
        "Newsletter", 
        "School Contact"
    ]

page = st.sidebar.selectbox("Go to", nav_options)

# --- HELPER FUNCTION FOR SUBJECT GRADING ($x < 1$ HANDLED AS NULL) ---
def get_subject_performance(score):
    if score is None or pd.isna(score) or float(score) < 1.0:
        return "-", "-", None
    score = float(score)
    if score <= 10: return "BE2", "Below Expectation 2", 1
    elif score <= 20: return "BE1", "Below Expectation 1", 2
    elif score <= 30: return "AE2", "Approaching Expectation 2", 3
    elif score <= 40: return "AE1", "Approaching Expectation 1", 4
    elif score <= 56: return "ME2", "Meeting Expectation 2", 5
    elif score <= 73: return "ME1", "Meeting Expectation 1", 6
    elif score <= 88: return "EE2", "Exceeding Expectation 2", 7
    else: return "EE1", "Exceeding Expectation 1", 8

def calculate_total_grade(total_marks):
    if total_marks <= 112: return "BE2 (Below Expectation 2)"
    elif total_marks <= 224: return "BE1 (Below Expectation 1)"
    elif total_marks <= 336: return "AE2 (Approaching Expectation 2)"
    elif total_marks <= 449: return "AE1 (Approaching Expectation 1)"
    elif total_marks <= 560: return "ME2 (Meeting Expectation 2)"
    elif total_marks <= 672: return "ME1 (Meeting Expectation 1)"
    elif total_marks <= 785: return "EE2 (Exceeding Expectation 2)"
    else: return "EE1 (Exceeding Expectation 1)"

# --- HELPER FUNCTION: FORMULATE TEACHER'S COMMENT ---
def generate_teacher_comment(total_score, top_subject, low_subject):
    overall_grade = calculate_total_grade(total_score)
    if "EE" in overall_grade:
        return f"An excellent performance! Demonstrates outstanding mastery across learning areas, especially in {top_subject}. Keep up the high standard!"
    elif "ME" in overall_grade:
        return f"Good performance. Shows consistent effort and understanding, particularly in {top_subject}. Aim to put more effort into {low_subject} for better results."
    elif "AE" in overall_grade:
        return f"Fair performance. Shows potential in {top_subject}, but requires more focus and practice in {low_subject} to improve overall standing."
    else:
        return f"Below expectations. Needs active academic support and dedicated revision, particularly in {low_subject}. Steady practice in {top_subject} will build confidence."

# --- PAGE 1: DASHBOARD ---
if page == "Dashboard":
    st.markdown("<h1 style='text-align: center;'>KEA COMPREHENSIVE SCHOOL</h1>", unsafe_allow_html=True)
    if os.path.exists("logo.png"):
        col_l1, col_l2, col_l3 = st.columns([1, 2, 1])
        with col_l2:
            st.image("logo.png", width=200)
    st.markdown("<h4 style='text-align: center; color: #555;'>Primary (Grade 1-6) & Junior Secondary (Grade 7-9)</h4>", unsafe_allow_html=True)
    st.markdown("---")

    total_students = 0
    total_teachers = 0
    total_collected = 0.0
    expected_fee_per_student = 550.0
    total_expected = 0.0
    deficit = 0.0

    if supabase:
        try:
            students_res = supabase.table("students").select("*", count="exact").execute()
            teachers_res = supabase.table("teachers").select("*", count="exact").execute()
            fees_res = supabase.table("fee_payments").select("amount_paid").execute()
            
            total_students = students_res.count if students_res.count is not None else len(students_res.data)
            total_teachers = teachers_res.count if teachers_res.count is not None else len(teachers_res.data)
            
            if fees_res.data:
                total_collected = sum(float(item.get("amount_paid", 0)) for item in fees_res.data)
                
            total_expected = total_students * expected_fee_per_student
            deficit = max(0.0, total_expected - total_collected)
        except Exception as e:
            st.error(f"Error fetching dashboard metrics: {e}")

    col1, col2, col3, col4 = st.columns(4)
    col1.metric("Total Students", total_students)
    col2.metric("Total Teachers", total_teachers)
    col3.metric("Fee Collected (Ksh)", f"{total_collected:,.2f}")
    col4.metric("Fee Deficit (Ksh)", f"{deficit:,.2f}", delta_color="inverse")

    st.markdown("### Financial Track & Overview")
    progress_val = (total_collected / total_expected) if total_expected > 0 else 0.0
    progress_val = min(1.0, max(0.0, float(progress_val)))
    
    st.progress(progress_val)
    st.info(f"Financial Progress: Ksh {total_collected:,.2f} collected out of expected Ksh {total_expected:,.2f} total termly revenue.")

# --- PAGE 2: STUDENTS REGISTRATION (RESTRICTED TO ADMIN) ---
elif page == "Students Registration":
    if not is_admin_role:
        st.error("Access Denied: Only HOI, DHOI, and Senior Teachers can register students.")
        st.stop()
        
    st.header("Students Registration Portal")
    tab_manual, tab_excel = st.tabs(["Manual Registration", "Excel Spreadsheet Upload"])
    
    with tab_manual:
        with st.form("manual_student_form"):
            adm_no = st.text_input("Admission Number")
            assessment_no = st.text_input("Assessment Number")
            name = st.text_input("Full Name")
            grade = st.selectbox("Grade", [f"Grade {i}" for i in range(1, 10)])
            submit_manual = st.form_submit_button("Admit Student")
            
            if submit_manual:
                if adm_no and name and grade:
                    try:
                        supabase.table("students").insert({
                            "adm_no": adm_no,
                            "assessment_no": assessment_no,
                            "name": name,
                            "grade": grade
                        }).execute()
                        st.success(f"Successfully admitted {name} to {grade}!")
                    except Exception as e:
                        st.error(f"Error saving student: {e}")
                else:
                    st.warning("Please fill in all mandatory fields.")

    with tab_excel:
        st.markdown("Upload an Excel spreadsheet with columns: **Adm No**, **Assessment No**, **Name**, **Grade**")
        uploaded_file = st.file_uploader("Choose Excel file", type=["xlsx", "xls"])
        selected_grade_bulk = st.selectbox("Target Grade for Bulk Upload", [f"Grade {i}" for i in range(1, 10)])
        
        if uploaded_file and st.button("Process & Admit Students"):
            try:
                df = pd.read_excel(uploaded_file)
                records = []
                for _, row in df.iterrows():
                    records.append({
                        "adm_no": str(row.get("Adm No")),
                        "assessment_no": str(row.get("Assessment No", "")),
                        "name": str(row.get("Name")),
                        "grade": str(row.get("Grade", selected_grade_bulk))
                    })
                supabase.table("students").upsert(records).execute()
                st.success(f"Successfully uploaded and admitted {len(records)} students!")
            except Exception as e:
                st.error(f"Error processing excel file: {e}")

# --- PAGE 3: MARKS ENTRY ---
elif page == "Marks Entry":
    st.header("Subject-Specific Marks Entry Portal")
    
    user_key = st.session_state.username
    teacher_data = TEACHER_ASSIGNMENTS.get(user_key)
    
    # 1. Determine Allowed Grades & Subjects based on Logged-in Teacher
    if is_admin_role or not teacher_data:
        allowed_grades = [f"Grade {i}" for i in range(1, 10)]
        teacher_grade_map = {g: LEARNING_AREAS for g in allowed_grades}
        st.info("Log-in Mode: Administrator Mode (Full Access)")
    else:
        teacher_grade_map = {}
        for item in teacher_data["assignments"]:
            teacher_grade_map[item["grade"]] = item["subjects"]
        allowed_grades = list(teacher_grade_map.keys())
        st.info(f"Welcome, {st.session_state.full_name}! Select an assigned grade and learning area below.")

    if not allowed_grades:
        st.warning("No grades currently assigned to your profile.")
        st.stop()

    # 2. Select Grade & Learning Area
    col_g, col_s = st.columns(2)
    with col_g:
        target_grade = st.selectbox("1. Select Grade", allowed_grades)
    
    assigned_subjects = teacher_grade_map.get(target_grade, LEARNING_AREAS)
    
    with col_s:
        target_subject = st.selectbox("2. Select Learning Area", assigned_subjects)

    st.markdown("---")
    st.subheader(f"Marking Sheet: {target_grade} — {target_subject}")
    st.caption("Enter scores out of 100%. Any score less than 1 (x < 1) is saved as Null (-) with no performance level. Previously saved marks are automatically loaded for editing.")

    # 3. Fetch Registered Students
    grade_num = target_grade.replace("Grade", "").strip()
    try:
        students_res = supabase.table("students").select("adm_no, name, grade").execute()
        all_students = students_res.data if students_res.data else []
        students = [
            s for s in all_students 
            if str(s.get("grade", "")).strip().lower() in [target_grade.lower(), grade_num, f"grade {grade_num}"]
        ]
        students = sorted(students, key=lambda x: str(x.get("adm_no", "")))
    except Exception as e:
        students = []
        st.error(f"Error fetching students from database: {e}")

    if not students:
        st.warning(f"No students found under '{target_grade}' in the database.")
    else:
        col_key = target_subject.lower().replace(" ", "_").replace(".", "").replace("(", "").replace(")", "")

        # 4. Fetch Existing Saved Marks Live from Database
        existing_marks_map = {}
        try:
            m_res = supabase.table("marks").select(f"adm_no, {col_key}").eq("grade", target_grade).execute()
            if m_res.data:
                for row in m_res.data:
                    existing_marks_map[row["adm_no"]] = row.get(col_key)
        except Exception as e:
            st.caption(f"Note: Could not load prior marks automatically: {e}")

        # 5. Build Data Table for Entry & Editing
        editor_data = []
        for s in students:
            adm = s["adm_no"]
            val = existing_marks_map.get(adm)
            # If mark exists and >= 1, pre-fill it live; otherwise 0.0
            score_val = float(val) if val is not None and float(val) >= 1.0 else 0.0
            
            editor_data.append({
                "Adm No": adm,
                "Student Name": s["name"],
                "Score (%)": score_val
            })

        df_editor = pd.DataFrame(editor_data)

        # Tabs for Manual Sheet & Excel
        tab_manual_marks, tab_excel_marks = st.tabs(["Manual Sheet Input", "Upload Excel Marks"])

        with tab_manual_marks:
            with st.form(key=f"form_{target_grade}_{col_key}"):
                st.markdown("### Class List Marks Input")
                edited_df = st.data_editor(
                    df_editor,
                    column_config={
                        "Adm No": st.column_config.TextColumn("Adm No", disabled=True),
                        "Student Name": st.column_config.TextColumn("Student Name", disabled=True),
                        "Score (%)": st.column_config.NumberColumn(
                            "Score (%)", 
                            min_value=0.0, 
                            max_value=100.0, 
                            step=1.0,
                            help="Enter score between 1 and 100. Scores < 1 are saved as Null (-)"
                        )
                    },
                    use_container_width=True,
                    hide_index=True
                )
                
                preview_submitted = st.form_submit_button("Preview Marks Sheet", type="secondary")

            # 6. PREVIEW AND FINAL CONFIRMATION
            if preview_submitted or st.session_state.get(f"preview_active_{target_grade}_{col_key}", False):
                st.session_state[f"preview_active_{target_grade}_{col_key}"] = True
                
                st.markdown("---")
                st.markdown("### 📋 Preview Marks & Performance Levels")
                
                preview_rows = []
                for _, row in edited_df.iterrows():
                    raw_score = row["Score (%)"]
                    grade_code, perf_level, pts = get_subject_performance(raw_score)
                    
                    preview_rows.append({
                        "Adm No": row["Adm No"],
                        "Student Name": row["Student Name"],
                        "Score (%)": raw_score if raw_score >= 1.0 else "-",
                        "Grade": grade_code,
                        "Performance Level": perf_level
                    })
                    
                preview_table = pd.DataFrame(preview_rows)
                st.dataframe(preview_table, use_container_width=True)

                if st.button("✅ Confirm & Submit to Results Analysis", type="primary"):
                    records_to_upsert = []
                    for _, row in edited_df.iterrows():
                        adm = str(row["Adm No"])
                        raw_score = row["Score (%)"]
                        
                        # Fetch existing row to preserve marks in OTHER subjects
                        existing = {}
                        try:
                            ex_res = supabase.table("marks").select("*").eq("adm_no", adm).execute()
                            if ex_res.data:
                                existing = ex_res.data[0]
                        except:
                            pass

                        rec = existing if existing else {"adm_no": adm, "grade": target_grade}
                        
                        # Enforce $x < 1$ as NULL
                        rec[col_key] = float(raw_score) if raw_score >= 1.0 else None
                        
                        # Recalculate Total Marks (summing non-null scores only)
                        total = 0.0
                        for sub in LEARNING_AREAS:
                            s_key = sub.lower().replace(" ", "_").replace(".", "").replace("(", "").replace(")", "")
                            val = rec.get(s_key)
                            if val is not None and float(val) >= 1.0:
                                total += float(val)
                        
                        rec["total_marks"] = total
                        records_to_upsert.append(rec)

                    try:
                        supabase.table("marks").upsert(records_to_upsert, on_conflict="adm_no").execute()
                        st.success(f"🎉 Successfully updated {target_subject} marks for {len(records_to_upsert)} students in {target_grade}!")
                        st.session_state[f"preview_active_{target_grade}_{col_key}"] = False
                    except Exception as e:
                        st.error(f"Failed to submit marks to database: {e}")

        with tab_excel_marks:
            st.markdown(f"Upload an Excel sheet for **{target_grade} — {target_subject}**.")
            st.caption("Ensure the file contains an **'Adm No'** column and a column named **'Score'** or matching the subject.")
            
            marks_file = st.file_uploader("Choose Excel File", type=["xlsx", "xls"], key=f"uploader_{target_grade}_{col_key}")
            
            if marks_file and st.button("Upload Excel Marks Sheet"):
                try:
                    df_upload = pd.read_excel(marks_file)
                    records_to_upsert = []
                    
                    score_col = None
                    for c in df_upload.columns:
                        if c.lower() in ["score", "marks", target_subject.lower()]:
                            score_col = c
                            break
                    
                    if not score_col and len(df_upload.columns) >= 2:
                        score_col = df_upload.columns[1]

                    for _, row in df_upload.iterrows():
                        adm = str(row["Adm No"]).strip()
                        raw_val = row[score_col]
                        
                        try:
                            score_num = float(raw_val)
                        except:
                            score_num = 0.0

                        existing = {}
                        try:
                            ex_res = supabase.table("marks").select("*").eq("adm_no", adm).execute()
                            if ex_res.data:
                                existing = ex_res.data[0]
                        except:
                            pass

                        rec = existing if existing else {"adm_no": adm, "grade": target_grade}
                        rec[col_key] = score_num if score_num >= 1.0 else None
                        
                        total = 0.0
                        for sub in LEARNING_AREAS:
                            s_key = sub.lower().replace(" ", "_").replace(".", "").replace("(", "").replace(")", "")
                            val = rec.get(s_key)
                            if val is not None and float(val) >= 1.0:
                                total += float(val)
                        rec["total_marks"] = total
                        records_to_upsert.append(rec)

                    supabase.table("marks").upsert(records_to_upsert, on_conflict="adm_no").execute()
                    st.success(f"Successfully processed and uploaded marks for {len(records_to_upsert)} students!")
                except Exception as e:
                    st.error(f"Error processing uploaded file: {e}")
                    
# --- PAGE 4: RESULTS ANALYSIS & REPORT FORMS HUB ---
elif page == "Results Analysis":
    st.header("Results Analysis Hub & Report Forms")
    
    analysis_grade = st.selectbox("Select Grade for Analysis", [f"Grade {i}" for i in range(1, 10)], key="analysis_grade_select")
    
    tab_overview, tab_reports = st.tabs(["Tab 1: General Performance Overview", "Tab 2: Assessment Report Forms Hub"])
    
    # 1. FETCH DATA WITH STRICT STRING NORMALIZATION
    grade_num = analysis_grade.replace("Grade", "").strip()
    try:
        m_res = supabase.table("marks").select("*").execute()
        all_marks = m_res.data if m_res.data else []
        marks_data = [
            m for m in all_marks 
            if str(m.get("grade", "")).strip().lower() in [analysis_grade.lower(), grade_num, f"grade {grade_num}"]
        ]
        
        s_res = supabase.table("students").select("*").execute()
        students_dict = {}
        if s_res.data:
            for s in s_res.data:
                clean_adm = str(s.get("adm_no", "")).strip()
                students_dict[clean_adm] = s
    except Exception as e:
        marks_data, students_dict = [], {}
        st.error(f"Error fetching analysis data: {e}")

    # HELPER TO CALCULATE VALID TOTAL & COUNT EXAMS TAKEN
    def process_student_scores(m_row):
        total = 0.0
        valid_count = 0
        for sub in LEARNING_AREAS:
            k = sub.lower().replace(" ", "_").replace(".", "").replace("(", "").replace(")", "")
            val = m_row.get(k)
            if val is not None and not pd.isna(val) and float(val) >= 1.0:
                total += float(val)
                valid_count += 1
        return total, valid_count

    # PREPARE ACTIVE RANKED STUDENTS LIST
    active_ranked_students = []
    for m in marks_data:
        adm_str = str(m.get("adm_no", "")).strip()
        s_name = students_dict.get(adm_str, {}).get("name", f"Student {adm_str}")
        tot_marks, valid_cnt = process_student_scores(m)
        
        if valid_cnt > 0:
            active_ranked_students.append({
                "adm_no": adm_str,
                "name": s_name,
                "total_marks": tot_marks,
                "marks_row": m
            })

    # Sort Active Students by Rank
    active_ranked_students = sorted(active_ranked_students, key=lambda x: x["total_marks"], reverse=True)
    for idx, item in enumerate(active_ranked_students):
        item["rank"] = idx + 1

    # =========================================================
    # TAB 1: GENERAL PERFORMANCE OVERVIEW
    # =========================================================
    with tab_overview:
        if not active_ranked_students:
            st.info(f"No active student records with submitted marks found in {analysis_grade}.")
        else:
            st.markdown("### 🏆 Top Ten Students")
            top_10_rows = []
            for item in active_ranked_students[:10]:
                top_10_rows.append({
                    "Rank": item["rank"],
                    "Adm No": item["adm_no"],
                    "Student Name": item["name"],
                    "Total Marks": item["total_marks"],
                    "Overall Grade": calculate_total_grade(item["total_marks"])
                })
            st.dataframe(pd.DataFrame(top_10_rows), use_container_width=True)
            
            st.markdown("---")

            subject_stats = []
            df_marks = pd.DataFrame(marks_data)
            
            for sub in LEARNING_AREAS:
                col_name = sub.lower().replace(" ", "_").replace(".", "").replace("(", "").replace(")", "")
                if col_name in df_marks.columns:
                    valid_scores = [float(v) for v in df_marks[col_name] if v is not None and not pd.isna(v) and float(v) >= 1.0]
                    valid_count = len(valid_scores)
                    
                    if valid_count > 0:
                        total_sub_marks = sum(valid_scores)
                        mean_score = total_sub_marks / valid_count
                        total_points = sum([get_subject_performance(s)[2] for s in valid_scores if get_subject_performance(s)[2] is not None])
                        mean_points = total_points / valid_count
                    else:
                        total_sub_marks, mean_score, mean_points = 0.0, 0.0, 0.0

                    subject_stats.append({
                        "Learning Area": sub,
                        "Valid Students": valid_count,
                        "Total Marks": round(total_sub_marks, 1),
                        "Mean Score (%)": round(mean_score, 2),
                        "Aggregate Mean Points": round(mean_points, 2)
                    })

            df_subject_stats = pd.DataFrame(subject_stats)

            if not df_subject_stats.empty and df_subject_stats["Mean Score (%)"].max() > 0:
                best_subject = df_subject_stats.sort_values(by="Mean Score (%)", ascending=False).iloc[0]
                st.markdown(f"### ⭐ Best Performed Learning Area")
                st.success(
                    f"**{best_subject['Learning Area']}** with a Mean Score of **{best_subject['Mean Score (%)']}%** "
                    f"and Aggregate Points of **{best_subject['Aggregate Mean Points']}** (based on {best_subject['Valid Students']} valid submissions)."
                )
            
            st.markdown("---")
            st.markdown("### 📊 Performance Summary for Each Learning Area")
            st.dataframe(df_subject_stats, use_container_width=True)

    # =========================================================
    # TAB 2: ASSESSMENT REPORT FORMS HUB
    # =========================================================
    with tab_reports:
        st.markdown("### Assessment Report Form Generator")
        
        col_s1, col_s2 = st.columns(2)
        with col_s1:
            term_val = st.selectbox("Term", ["Term 1", "Term 2", "Term 3"], key="rep_term")
            opening_date = st.text_input("Opening Date", "14/09/2026", key="rep_open")
        with col_s2:
            closing_date = st.text_input("Closing Date", "04/12/2026", key="rep_close")
            stamp_upload = st.file_uploader("Upload Stamp Image (stamp.png)", type=["png", "jpg"], key="rep_stamp")

        st.markdown("---")
        
        if not active_ranked_students:
            st.info(f"No students with active exam scores found in {analysis_grade}. Report forms are suppressed for students without marks.")
        else:
            # HELPER FUNCTION TO POPULATE WORD DOCUMENT TEMPLATE EXACTLY
            def generate_student_docx_object(student_obj):
                adm_str = student_obj["adm_no"]
                student_name = student_obj["name"]
                s_row = student_obj["marks_row"]
                student_rank = student_obj["rank"]

                if os.path.exists("report.docx"):
                    doc = Document("report.docx")
                else:
                    doc = Document()

                total_score = 0.0
                top_sub, low_sub = "Mathematics", "English"
                best_s, worst_s = -1, 101

                # Calculate marks dictionary
                subject_scores = {}
                for sub in LEARNING_AREAS:
                    col_key = sub.lower().replace(" ", "_").replace(".", "").replace("(", "").replace(")", "")
                    val = s_row.get(col_key)
                    if val is not None and not pd.isna(val) and float(val) >= 1.0:
                        s_num = float(val)
                        total_score += s_num
                        if s_num > best_s: best_s = s_num; top_sub = sub
                        if s_num < worst_s: worst_s = s_num; low_sub = sub
                        g_code, p_lvl, pts = get_subject_performance(s_num)
                        subject_scores[sub.upper()] = (str(s_num), p_lvl)
                    else:
                        subject_scores[sub.upper()] = ("-", "-")

                overall_g = calculate_total_grade(total_score)
                teacher_comment = generate_teacher_comment(total_score, top_sub, low_sub)

                # 1. Update Header / Info Tables in Word Document
                for t in doc.tables:
                    for row in t.rows:
                        row_text = " ".join([cell.text for cell in row.cells]).upper()
                        
                        # Populate Student Header Details
                        for cell in row.cells:
                            if "NAME:" in cell.text.upper():
                                cell.text = f"NAME: {student_name}"
                            elif "ADM NO:" in cell.text.upper() or "ADM.NO:" in cell.text.upper():
                                cell.text = f"ADM NO: {adm_str}"
                            elif "TERM:" in cell.text.upper():
                                cell.text = f"TERM: {term_val}"
                            elif "POSITION:" in cell.text.upper() or "RANK:" in cell.text.upper():
                                cell.text = f"POSITION: #{student_rank}"
                            elif "YEAR:" in cell.text.upper():
                                cell.text = f"YEAR: 2026"
                                
                        # Populate Learning Area Table Rows Directly
                        for sub_name, (score_str, perf_str) in subject_scores.items():
                            if sub_name in row_text and len(row.cells) >= 5:
                                row.cells[3].text = score_str
                                row.cells[4].text = perf_str

                # 2. Update Paragraphs (Total Marks, Comments, Signatures at bottom)
                replacements = {
                    "TOTAL MARKS: _____": f"TOTAL MARKS: {total_score:.0f}",
                    "GENERAL PERFORMANCE LEVEL:": f"GENERAL PERFORMANCE LEVEL: {overall_g}",
                    "TEACHERS GENERAL COMMENT:": f"TEACHERS GENERAL COMMENT:\n\"{teacher_comment}\"",
                    "CLOSING DATE: __________": f"CLOSING DATE: {closing_date}",
                    "OPENING DATE: __________": f"OPENING DATE: {opening_date}",
                }

                for p in doc.paragraphs:
                    for k, v in replacements.items():
                        if k in p.text:
                            p.text = p.text.replace(k, v)

                # Insert Stamp
                if stamp_upload or os.path.exists("stamp.png"):
                    stamp_source = stamp_upload if stamp_upload else "stamp.png"
                    for p in doc.paragraphs:
                        if "H.O.I STAMP" in p.text or "{{STAMP}}" in p.text:
                            run = p.add_run()
                            run.add_picture(stamp_source, width=Inches(1.2))
                            break

                return doc

            # --- OPTION 1: LIVE PREVIEW & SINGLE REPORT DOWNLOAD ---
            st.subheader("1. Live Student Report Form Preview")
            
            selected_student_obj = st.selectbox(
                "Select Student for Live Report Form Preview", 
                active_ranked_students, 
                format_func=lambda x: f"Rank #{x['rank']} | Adm No: {x['adm_no']} | Name: {x['name']} (Total: {x['total_marks']:.0f})"
            )

            if selected_student_obj:
                doc_obj = generate_student_docx_object(selected_student_obj)
                
                s_name = selected_student_obj['name']
                s_adm = selected_student_obj['adm_no']
                s_rank = selected_student_obj['rank']
                s_total = selected_student_obj['total_marks']
                s_row = selected_student_obj['marks_row']
                overall_g = calculate_total_grade(s_total)
                
                # Derive Comments
                top_s, low_s = "Mathematics", "English"
                b_val, w_val = -1, 101
                for sub in LEARNING_AREAS:
                    col_key = sub.lower().replace(" ", "_").replace(".", "").replace("(", "").replace(")", "")
                    val = s_row.get(col_key)
                    if val is not None and not pd.isna(val) and float(val) >= 1.0:
                        v = float(val)
                        if v > b_val: b_val = v; top_s = sub
                        if v < w_val: w_val = v; low_s = sub
                teacher_comment = generate_teacher_comment(s_total, top_s, low_s)

                st.markdown("#### 👁️ Live Report Form Document View")
                
                # STRUCTURED HTML rendering: Header -> Info -> Learning Area Table -> Footer Summaries
                html_preview = f"""
                <div style="border: 2px solid #374151; padding: 25px; border-radius: 8px; background-color: #FFFFFF; color: #111827; font-family: Arial, sans-serif;">
                    
                    <!-- HEADER -->
                    <div style="text-align: center; font-weight: bold; margin-bottom: 15px;">
                        <h2 style="margin: 0; color: #1E3A8A;">KEA COMPREHENSIVE SCHOOL</h2>
                        <p style="margin: 2px 0;">P.O. BOX 557-40400 SUNA MIGORI</p>
                        <h4 style="margin: 5px 0; text-decoration: underline;">STUDENT ASSESSMENT REPORT</h4>
                    </div>

                    <!-- STUDENT DETAILS TABLE -->
                    <table style="width: 100%; border-collapse: collapse; margin-bottom: 15px; font-size: 13px; border: 1px solid #374151;">
                        <tr>
                            <td style="border: 1px solid #374151; padding: 6px; font-weight: bold;">NAME: {s_name}</td>
                            <td style="border: 1px solid #374151; padding: 6px; font-weight: bold;">YEAR: 2026</td>
                        </tr>
                        <tr>
                            <td style="border: 1px solid #374151; padding: 6px; font-weight: bold;">TERM: {term_val}</td>
                            <td style="border: 1px solid #374151; padding: 6px; font-weight: bold;">POSITION: #{s_rank}</td>
                        </tr>
                        <tr>
                            <td style="border: 1px solid #374151; padding: 6px; font-weight: bold;">ADM NO: {s_adm}</td>
                            <td style="border: 1px solid #374151; padding: 6px; font-weight: bold;">GRADE: {analysis_grade}</td>
                        </tr>
                    </table>

                    <!-- LEARNING AREAS MARKS TABLE -->
                    <table style="width: 100%; border-collapse: collapse; font-size: 12px; border: 1px solid #374151;">
                        <tr style="background-color: #F3F4F6; font-weight: bold; text-align: left;">
                            <th style="border: 1px solid #374151; padding: 6px;">S/N</th>
                            <th style="border: 1px solid #374151; padding: 6px;">CODE</th>
                            <th style="border: 1px solid #374151; padding: 6px;">LEARNING AREA</th>
                            <th style="border: 1px solid #374151; padding: 6px;">MARKS SCORED</th>
                            <th style="border: 1px solid #374151; padding: 6px;">PERFORMANCE LEVEL</th>
                        </tr>
                """

                # Populate table rows dynamically for each subject
                for idx, sub in enumerate(LEARNING_AREAS, 1):
                    col_key = sub.lower().replace(" ", "_").replace(".", "").replace("(", "").replace(")", "")
                    val = s_row.get(col_key)
                    
                    code_val = f"90{idx}" if idx < 10 else f"9{idx}"
                    
                    if val is not None and not pd.isna(val) and float(val) >= 1.0:
                        s_num = float(val)
                        g_code, p_lvl, pts = get_subject_performance(s_num)
                        m_str = f"{s_num:.0f}"
                        p_str = p_lvl
                    else:
                        m_str = "-"
                        p_str = "-"

                    html_preview += f"""
                        <tr>
                            <td style="border: 1px solid #374151; padding: 6px;">{idx}</td>
                            <td style="border: 1px solid #374151; padding: 6px;">{code_val}</td>
                            <td style="border: 1px solid #374151; padding: 6px; font-weight: bold;">{sub.upper()}</td>
                            <td style="border: 1px solid #374151; padding: 6px; font-weight: bold; color: #1E3A8A;">{m_str}</td>
                            <td style="border: 1px solid #374151; padding: 6px;">{p_str}</td>
                        </tr>
                    """

                # SUMMARY SECTION BELOW THE TABLE
                html_preview += f"""
                    </table>

                    <div style="margin-top: 15px; font-size: 13px; border-top: 2px solid #374151; padding-top: 10px;">
                        <p style="margin: 5px 0;"><strong>TOTAL MARKS:</strong> <span style="font-size: 15px; color: #1E3A8A;">{s_total:.0f}</span> &nbsp;|&nbsp; <strong>GENERAL PERFORMANCE LEVEL:</strong> {overall_g}</p>
                        
                        <p style="margin: 10px 0 3px 0;"><strong>TEACHER'S GENERAL COMMENT:</strong></p>
                        <div style="border-bottom: 1px dashed #374151; padding: 4px; font-style: italic; background-color: #F9FAFB;">"{teacher_comment}"</div>
                        
                        <div style="display: flex; justify-content: space-between; margin-top: 15px;">
                            <div><strong>CLOSING DATE:</strong> {closing_date}</div>
                            <div><strong>OPENING DATE:</strong> {opening_date}</div>
                        </div>

                        <div style="display: flex; justify-content: space-between; margin-top: 20px; font-weight: bold;">
                            <div>CLASS TEACHER SIGNATURE: ____________</div>
                            <div>H.O.I STAMP & SIGNATURE: ____________</div>
                        </div>
                    </div>

                </div>
                """

                st.components.v1.html(html_preview, height=620, scrolling=True)

                # Download Single File
                out_stream = io.BytesIO()
                doc_obj.save(out_stream)
                out_stream.seek(0)
                
                st.download_button(
                    label=f"📥 Download Assessment Report Form for {s_name} (.docx)",
                    data=out_stream,
                    file_name=f"Rank_{s_rank}_{s_adm}_{s_name.replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary"
                )

            st.markdown("---")

            # --- OPTION 2: RANKED BATCH REPORT ZIP GENERATOR ---
            st.subheader("2. Ranked Batch Reports Generator (Entire Grade)")
            st.caption(f"Generates report forms ordered from Rank #1 to Rank #{len(active_ranked_students)}. Students without marks are excluded.")

            if st.button(f"📦 Generate All Ranked Reports for {analysis_grade} (.zip)", type="secondary"):
                with st.spinner("Compiling ranked report forms... Please wait."):
                    zip_buffer = io.BytesIO()
                    
                    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
                        for s_item in active_ranked_students:
                            single_doc = generate_student_docx_object(s_item)
                            
                            doc_bytes = io.BytesIO()
                            single_doc.save(doc_bytes)
                            doc_bytes.seek(0)
                            
                            file_filename = f"Rank_{s_item['rank']:02d}_{s_item['adm_no']}_{s_item['name'].replace(' ', '_')}_Report.docx"
                            zip_file.writestr(file_filename, doc_bytes.getvalue())

                    zip_buffer.seek(0)
                    st.success(f"Successfully compiled {len(active_ranked_students)} ranked report forms into ZIP archive!")
                    
                    st.download_button(
                        label=f"📥 Download All {len(active_ranked_students)} Ranked Reports (.zip)",
                        data=zip_buffer,
                        file_name=f"{analysis_grade.replace(' ', '_')}_Ranked_Assessment_Reports.zip",
                        mime="application/zip"
                )
                
# --- PAGE 5: FEE PAYMENT PORTAL (ADMIN ONLY) ---
elif page == "Fee Payment":
    if not is_admin_role:
        st.error("Access Denied: Only HOI, DHOI, and Senior Teachers have access to Financial Records.")
        st.stop()
        
    st.header("Fee Payment & Financial Portal")
    tab_fee_struct, tab_make_pay, tab_history, tab_balances = st.tabs(["Fee Structure", "Make Payment", "Receipt & History", "Termly Balances"])
    
    with tab_fee_struct:
        st.markdown("### KEA Comprehensive School Fee Structure")
        st.info("""
        * **Continuing Students:** Ksh 550 per term (Tuition: Ksh 150 [50/month], Exams: Ksh 400 [200 Mid Term, 200 End Term]). Total Annual = Ksh 1,650.
        * **Newly Admitted Students:** Pay an extra Ksh 200 admission fee in Term 1 (Total Term 1 = Ksh 750). Terms 2 & 3 revert to standard Ksh 550.
        """)
        
    with tab_make_pay:
        try:
            st_res = supabase.table("students").select("*").execute()
            all_students = st_res.data
        except:
            all_students = []
            
        if all_students:
            with st.form("payment_form"):
                payer_student = st.selectbox("Search / Select Student", all_students, format_func=lambda x: f"{x['adm_no']} - {x['name']} ({x['grade']})")
                amount = st.number_input("Amount Paid (Ksh)", min_value=0.0, value=550.0)
                channel = st.selectbox("Payment Channel", ["Cash Payment", "Bank Payment", "M-Pesa Payment"])
                payer_name = st.text_input("Name of Payer")
                category = st.selectbox("Payment Category", ["Whole Term Payment", "Tuition Fee", "Exams Fee", "Admission Fee"])
                
                submit_pay = st.form_submit_button("Record Payment")
                if submit_pay:
                    try:
                        supabase.table("fee_payments").insert({
                            "adm_no": payer_student["adm_no"],
                            "student_name": payer_student["name"],
                            "grade": payer_student["grade"],
                            "amount_paid": amount,
                            "channel": channel,
                            "payer_name": payer_name,
                            "category": category,
                            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                        }).execute()
                        st.success(f"Successfully recorded payment of Ksh {amount} for {payer_student['name']}!")
                    except Exception as e:
                        st.error(f"Error recording payment: {e}")

    with tab_history:
        st.markdown("### Payment History & Receipt Generation")
        try:
            pay_res = supabase.table("fee_payments").select("*").order("timestamp", desc=True).execute()
            payments = pay_res.data
            if payments:
                df_pay = pd.DataFrame(payments)
                st.dataframe(df_pay, use_container_width=True)
            else:
                st.info("No payment history recorded yet.")
        except Exception as e:
            st.error(f"Could not load payment history: {e}")

    with tab_balances:
        st.markdown("### Termly Balances Status")
        try:
            st_res = supabase.table("students").select("*").execute()
            pay_res = supabase.table("fee_payments").select("*").execute()
            
            student_paid = {}
            for p in pay_res.data:
                adm = p["adm_no"]
                student_paid[adm] = student_paid.get(adm, 0) + p["amount_paid"]
                
            balance_rows = []
            for s in st_res.data:
                paid = student_paid.get(s["adm_no"], 0)
                required = 550
                diff = paid - required
                if diff == 0:
                    status = "🟢 Nil Balance"
                elif diff < 0:
                    status = f"🔴 Pending Balance (Deficit: {-diff})"
                else:
                    status = f"🔵 Overpayment (Credit: +{diff})"
                    
                balance_rows.append({
                    "Adm No": s["adm_no"],
                    "Name": s["name"],
                    "Grade": s["grade"],
                    "Paid (Ksh)": paid,
                    "Status": status
                })
            
            if balance_rows:
                st.dataframe(pd.DataFrame(balance_rows), use_container_width=True)
        except Exception as e:
            st.error(f"Error calculating balances: {e}")

 # --- PAGE 6: TEACHERS PORTAL ---
elif page == "Teachers Portal":
    st.header(f"Teacher Portal — {st.session_state.full_name}")
    
    user_key = st.session_state.username
    teacher_info = TEACHER_ASSIGNMENTS.get(user_key)
    
    if is_admin_role and not teacher_info:
        st.markdown("### Administrator Overview")
        st.info("You are logged in with Administrative privileges.")
        try:
            t_res = supabase.table("teachers").select("full_name", "username", "designation").execute()
            st.dataframe(pd.DataFrame(t_res.data), use_container_width=True)
        except:
            pass
    elif teacher_info:
        st.markdown("### Your Assigned Learning Areas & Responsibilities")
        
        # Display ONLY the logged in teacher's assignments
        for item in teacher_info["assignments"]:
            grade = item["grade"]
            subs = ", ".join(item["subjects"])
            st.success(f"📌 **{grade}:** Teaching **{subs}**")
            
            # List students in this teacher's assigned class directly on their portal
            try:
                st_res = supabase.table("students").select("adm_no", "assessment_no", "name").eq("grade", grade).execute()
                if st_res.data:
                    st.markdown(f"**Registered Students in {grade}:**")
                    st.dataframe(pd.DataFrame(st_res.data), use_container_width=True)
                else:
                    st.caption(f"No students currently registered under {grade}.")
            except Exception as e:
                st.error(f"Could not load students for {grade}: {e}")
            st.markdown("---")
    else:
        st.warning("No teaching assignments found for your account.")

# --- PAGE 7: TEACHER TIME LOGIN ---
elif page == "Teacher Time Login":
    st.header("Teacher Attendance & Time Log")
    
    col1, col2 = st.columns(2)
    with col1:
        if st.button("Clock In (Time In)", type="primary"):
            st.session_state.show_in_popup = True
            
        if st.session_state.get("show_in_popup", False):
            with st.form("time_in_form"):
                t_name = st.text_input("Enter Your First Name")
                if st.form_submit_button("Submit Time In"):
                    try:
                        supabase.table("teacher_time_logs").insert({
                            "teacher_name": t_name,
                            "time_in": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                            "day_of_week": datetime.now().strftime("%A")
                        }).execute()
                        st.success(f"Successfully clocked IN for {t_name}!")
                        st.session_state.show_in_popup = False
                    except Exception as e:
                        st.error(f"Error recording time in: {e}")

    with col2:
        if st.button("Clock Out (Time Out)", type="secondary"):
            st.session_state.show_out_popup = True
            
        if st.session_state.get("show_out_popup", False):
            with st.form("time_out_form"):
                t_name_out = st.text_input("Enter Your First Name for Clock Out")
                if st.form_submit_button("Submit Time Out"):
                    try:
                        supabase.table("teacher_time_logs").update({
                            "time_out": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                        }).eq("teacher_name", t_name_out).is_("time_out", "null").execute()
                        st.success(f"Successfully clocked OUT for {t_name_out}!")
                        st.session_state.show_out_popup = False
                    except Exception as e:
                        st.error(f"Error recording time out: {e}")

    st.markdown("### Weekly Attendance Logs (Mon - Fri)")
    try:
        logs_res = supabase.table("teacher_time_logs").select("*").execute()
        if logs_res.data:
            st.dataframe(pd.DataFrame(logs_res.data), use_container_width=True)
    except:
        st.info("No attendance logs found.")

# --- PAGE 8: NEWSLETTER ---
elif page == "Newsletter":
    st.header("School Newsletter & Upcoming Events")
    
    if is_admin_role:
        with st.form("newsletter_form"):
            news_title = st.text_input("Newsletter Title")
            news_content = st.text_area("News Content & Announcements")
            submit_news = st.form_submit_button("Publish Newsletter")
            if submit_news:
                try:
                    supabase.table("newsletter").insert({
                        "title": news_title,
                        "content": news_content,
                        "date": datetime.now().strftime("%Y-%m-%d")
                    }).execute()
                    st.success("Newsletter published successfully!")
                except Exception as e:
                    st.error(f"Error publishing: {e}")
    else:
        st.info("Note: You are viewing in read-only mode. Only administrators can publish updates.")

    try:
        n_res = supabase.table("newsletter").select("*").order("date", desc=True).execute()
        for item in n_res.data:
            st.markdown(f"### {item['title']}")
            st.markdown(f"*{item['date']}*")
            st.write(item['content'])
            st.markdown("---")
    except:
        st.info("No newsletters posted yet.")

# --- PAGE 9: SCHOOL CONTACT ---
elif page == "School Contact":
    st.header("School Contact & Location Information")
    
    if is_admin_role:
        with st.form("contact_form"):
            address = st.text_input("Physical Address", "Migori County, Kenya")
            phone = st.text_input("School Contact Number", "+254 700 000 000")
            email = st.text_input("Email Address", "info@keacomprehensiveschool.ac.ke")
            maps_link = st.text_input("Google Maps Embed URL / Location Link", "https://maps.google.com")
            
            if st.form_submit_button("Update Contact Info"):
                try:
                    supabase.table("school_contact").upsert({
                        "id": 1,
                        "address": address,
                        "phone": phone,
                        "email": email,
                        "maps_link": maps_link
                    }).execute()
                    st.success("Contact information updated!")
                except Exception as e:
                    st.error(f"Error updating contact: {e}")
    
    st.markdown("### Contact Details")
    st.markdown("**Physical Address:** Migori County, Kenya")
    st.markdown("**Phone:** +254 700 000 000")
    st.markdown("**Email:** info@keacomprehensiveschool.ac.ke")
    st.markdown("[View Location on Google Maps](https://maps.google.com)")
